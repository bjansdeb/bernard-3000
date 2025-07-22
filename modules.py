import xlsxwriter
import openpyxl
from openpyxl import Workbook
from openpyxl import load_workbook
from urllib.request import urlopen
from bs4 import BeautifulSoup
from tkinter import *
from tkinter.filedialog import askopenfilename
import tkinter.messagebox
import docx
import re
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import RGBColor

def choosing_file():
    global c_file
    c_file = askopenfilename(filetypes=[("docx files", "*.docx")],initialdir="/",title='Choisir un fichier à préparer')
    if len(c_file) > 0:
        tkinter.messagebox.showinfo("Fichier ajouté !", "Bernard a bien reçu votre fichier! \nVous avez choisi le répertoire %s" %c_file)

def preparator():
##  try :
  document = docx.Document(c_file)
  #création des styles dans le document word
  style = document.styles
  style_names = [styl.name for styl in style] #Reprend tous les styles de docx pour vérifier si on les a

  if 'Cote' not in style_names:
      cote = style.add_style('Cote', WD_STYLE_TYPE.PARAGRAPH)
     
  else:
      cote = style['Cote']
  if 'Description' not in style_names:
      description = style.add_style('Description', WD_STYLE_TYPE.PARAGRAPH)
     
  else:
      description = style['Description']
  if 'Datum' not in style_names:
      datum = style.add_style('Datum', WD_STYLE_TYPE.PARAGRAPH)
  else:
      datum = style['Datum']
     
  if 'Importance' not in style_names:
      importance = style.add_style('Importance', WD_STYLE_TYPE.PARAGRAPH)
     
  else:
      importance = style['Importance']

  if 'Note' not in style_names:
      note = style.add_style('Note', WD_STYLE_TYPE.PARAGRAPH)
  else:
      note = style['Note']

  if 'TitreArch' not in style_names:
      t_arch = style.add_style('TitreArch', WD_STYLE_TYPE.PARAGRAPH)
     
  else:
      t_arch = style['TitreArch']
     
  cote.font.color.rgb = RGBColor(230, 4, 34) #Rouge Ajoute une couleur aux nouveaux styles
  description.font.color.rgb = RGBColor(127, 142, 228) #Bleu
  datum.font.color.rgb = RGBColor(51, 204, 102) #Vert
  importance.font.color.rgb = RGBColor(148, 51, 205) #Violet
  note.font.color.rgb = RGBColor(205, 169, 51) #Orange
  t_arch.font.color.rgb = RGBColor(100, 79, 81) #Brun/gris

  ident = re.compile(r"([1-9]\.|[1-9]\’\.|[1-9][0-9]\.|[1-9][0-9]\’\.|[1-9][0-9][0-9]\.|[1-9][0-9][0-9]\’\.|[0-9]\-[0-9]\.|[0-9]\-[0-9]\s:|[0-9]\-[0-9][0-9]\.|[0-9]\-[0-9][0-9]\.|[1-9][0-9]{3}\.|[1-9][0-9]{3}\’\.)")
  t = re.compile(r"[0-9]\.\\t([A-Za-z\s\'\ô\ö\ç\â\à\ä\ê\é\è\ë\ï\ü\-\—\’\'\(\)\,\“\”\«\»\;\:\\\\&\[\]\°\?\.]+)\\t") #|[0-9]\\t([A-Za-z\s\'\ô\ö\ç\â\à\ä\ê\é\è\ë\ï\ü\-\—\’\'\(\)\,\“\”\«\»\;\:\\\\&\[\]\°\?\.]+)\\t|[0-9]\\t([A-Za-z\s\'\ô\ö\ç\â\à\ä\ê\é\è\ë\ï\ü\-\—\’\'\(\)\,\“\”\«\»\;\:\\\\&\[\]\°\?\.]+)"
  t_1 = re.compile(r"([A-Za-z\s\'\ô\ö\ç\â\à\ä\ê\é\è\ë\ï\ü\-\—\’\'\(\)\,\“\”\«\»\;\:\\\\&\[\]\°\?\.]+)")
  imp = re.compile(r"^1 p...e|1 p...es|[0-9] p...es|[1-9][0-9] p...es|1 ch....e|1 ch....es|[0-9] ch....es|1 li...e|1 li...es|[0-9] li....s|1 vol..e|[1-9] vol..es$")
  d = re.compile (r"[1-9][0-9][0-9][0-9]-[0-9][0-9]|20[0-9][0-9]|[1-9][0-9][0-9][0-9]|.\.d\.|sd\.|sans date")
  date = re.compile (r"(?<![a-z]\s|:\s|\S0)[1-9][7-9][0-9]{2}-[1-9][7-9][0-9]{2}|(?<![a-z]\s|\,\s)[1-9][7-9][0-9]{2}|(?<![a-z]\s|\,\s)20[0-9]{2}|.\.d\.|sd\.|sans date|\[[1-9][7-9][0-9]{2}-[1-9][7-9][0-9]{2}\]|[ca\.[1-9][7-9][0-9]{2}]")
  #tabulation = re.compile(r"\t")
  #number_point = re.compile(r"[0-9]\.")
  for paragraph in document.paragraphs:
    #for run in runs
    if '\t' in paragraph.text:
      lines = paragraph.text.split('\t')
      paragraph.text = lines[0]
      for line in lines[1:]:
        new_paragraph = document.add_paragraph()
        new_paragraph.text = line

       
##      paragraph.text = re.sub(tabulation, "\n", paragraph.text)
##      paragraph.insert_paragraph_before()
     
##    if imp.search(paragraph.text):
##      text_copy = paragraph.text
##      text_s = re.sub(imp,"", text_copy)
##      paragraph.insert_paragraph_before(text_s)
##      paragraph.text = re.sub(d,"" , paragraph.text)
     
   
  for paragraph in document.paragraphs:
    if ident.search(paragraph.text):
      paragraph.style = cote
    if t.search(paragraph.text):
      paragraph.style = description
    if t_1.search(paragraph.text):
      paragraph.style = description
    if date.search(paragraph.text):
      paragraph.style = datum
    if imp.search(paragraph.text):
      paragraph.style = importance
    for run in paragraph.runs :
      if run.font.bold == True :
        paragraph.style = t_arch
      if run.font.all_caps == True or run.font.small_caps == True :
        paragraph.style = t_arch

  document.save("Bernard_auto.docx")
##    tkinter.messagebox.showinfo("Succès!", "La préparation a réussi, le résultat est visible dans le répertoire du logiciel \nIl porte le nom 'Bernard_auto'")
##  except:
##    tkinter.messagebox.showinfo("Echec!", "La préparation est un échec \nIl se peut qu'un fichier Bernard_auto soit déjà ouvert ou que vous n'ayez pas choisi de fichier")

def choixRep():
  global rep
  rep = askopenfilename(filetypes=[("htm files","*.htm"), ("html files","*.html")],initialdir="/",title='Choisir un fichier')
  global html
  if len(rep) > 0:
    tkinter.messagebox.showinfo("Fichier ajouté !", "Bernard a bien reçu votre fichier! \nVous avez choisi le répertoire %s" % rep)
    html = urlopen('file:///%s' % rep)

def extractor():
  try:
   bs = BeautifulSoup(html, 'html.parser')
   #creation excel et workbook
   myFile = xlsxwriter.Workbook('Bernard.xlsx')
   worksheet = myFile.add_worksheet()
   headings = ['Numéro', 'Titre unité', 'Dates', 'Importance', 'Notes']
   worksheet.write_row('A1', headings)
   #variables
   row = 0
   col = 0
   row2= 0
   row3 = 0
   row4 = 0
   row5 = 0

   #nettoyage du fichier htm avant extraction
   for x in bs.find_all():
     if len(x.get_text(strip=True)) == 0:
      x.extract()
   #print(bs)
   for tag in bs.find_all('p'):
       #Nettoyeur de balises 'o:p'
       check = tag.find('o:p')
       if(check != None):
           #check.decompose() était utilisé auparavant
           check.replace_with(" ")
   #print(bs)

   for name in bs.find_all('p'):
     if name.has_attr('class'):
       if name['class'][0] == 'Cote':
         row+=1
         worksheet.write(row, col, str(name.get_text()))
          ##print(name)
       if name['class'][0] == 'Description':
         row2=row
         worksheet.write(row2, col+1, str(name.get_text()))
       if name['class'][0] == 'Datum':
         row3=row
         worksheet.write(row3,col+2, str(name.get_text()))
       if name['class'][0] == 'Importance':
         row4=row
         worksheet.write(row4,col+3, str(name.get_text()))
       if name['class'][0] == 'Note':
         row5=row
         worksheet.write(row5,col+4, str(name.get_text()))
       if name['class'][0] == 'TitreArch':
         row+=1
         row2=row
         worksheet.write(row2, col+1, str(name.get_text()))
         

   myFile.close()
   tkinter.messagebox.showinfo("Succès!", "L'extraction a réussi. \nDonnées traitées %s" %row)
  except:
    tkinter.messagebox.showerror("Erreur", "Impossible d'extraire les données. \nVeuillez d'abord choisir un fichier ou fermer le fichier Bernard.xlxs!")

def choixEpurator():
  global epur
  epur = askopenfilename(filetypes=[("xlsx files","*.xlsx")],initialdir="/",title='Choisir un fichier à épurer')
  global exEpur
  if len(epur) > 0:
    tkinter.messagebox.showinfo("Fichier ajouté !", "Bernard a bien reçu votre fichier! \nVous avez choisi le repertoire %s" %epur)
    exEpur = load_workbook(epur)
   
def epurator():
  """Nettoyage de données avant import ; je l'ai fait de deux manières pour être sûr"""
  try :
    ws = exEpur.active
    pattern = r"[\x00-\x1F\x7F-\x9F\@\n\[\]\$\&]"
    for row in ws:
      for cell in row :
        if(cell.value != None):
          value = str(cell.value).strip()
          value_1st_round = re.sub(pattern,' ', value)
          # cell.value = value_1st_round.replace('_x000D_','').replace('@','').replace('\n',' ').replace(' ',' ').replace('[', '').replace(']','').replace('$','').replace('&','').replace('<0xa0>','').replace(' - ','\n-')
          cell.value = value_1st_round.replace('_x000D_','').replace('<0xa0>','').replace(' - ','\n-').replace('  ',' ').replace(' ;', ';').replace(' ','')
          ##replacements = { "_x000D_": "", "@": "", " ":" ", "\n":" "}
      
    exEpur.save('Epure.xlsx')
    tkinter.messagebox.showinfo("Success!", "L'épuration a réussi, le résultat est visible dans le répertoire du logiciel \nIl porte le nom 'Epure'")
  except:
    tkinter.messagebox.showerror("Erreur", "Impossible d'épurer les données. \nIl se peut que le fichier d'épuration soit déjà ouvert ou que vous n'ayez pas choisi de fichier")

def assign_at():
  """Transforme le template d'export en template d'import AtoM"""
  try:
    bernard = openpyxl.load_workbook('Epure.xlsx')
    assign = Workbook()
    bws = bernard.active
    aws = assign.active

    for i in range(2 ,bws.max_row +1):
      identifier = bws.cell(row=i, column= 1).value
      description = bws.cell(row=i, column= 2).value
      dates = bws.cell(row=i, column= 3).value
      importance = bws.cell(row=i,column= 4).value
      notes = bws.cell(row=i, column= 5).value

      aws.cell(i, 5).value = identifier
      aws.cell(i, 6).value = description
      aws.cell(i, 50).value = dates
      aws.cell(i, 8).value = importance
      aws.cell(i, 29).value = notes
      aws.cell(i, 56).value = 'fr'

    headers = ['legacyId','parentId','qubitParentSlug','accessionNumber', 'identifier','title', 'levelOfDescription', 'extentAndMedium', 'repository','archivalHistory', 'acquisition', 'scopeAndContent','appraisal','accruals','arrangement', 'accessConditions', 'reproductionConditions','language', 'script','languageNote', 'physicalCharacteristics', 'findingAids', 'locationOfOriginals','locationOfCopies', 'relatedUnitsOfDescription','publicationNote', 'digitalObjectPath', 'digitalObjectURI', 'generalNote', 'subjectAccessPoints', 'placeAccessPoints', 'nameAccessPoints', 'genreAccessPoints', 'descriptionIdentifier','institutionIdentifier', 'rules','descriptionStatus','levelOfDetail','revisionHistory','languageOfDescription', 'scriptOfDescription', 'sources', 'archivistNote','publicationStatus', 'physicalObjectName','physicalObjectLocation', 'physicalObjectType', 'alternativeIdentifiers', 'alternativeIdentifierLabels','eventDates','eventTypes','eventStartDates','eventEndDates','eventActors','eventActorHistories','culture']
    r = 1
    for names in headers:
      aws.cell(1, r).value = names
      r+=1

    assign.save('Template_bernard_atom.xlsx')
    tkinter.messagebox.showinfo("Success", "Bernard a tout swappé en template atom \n Le fichier a été sauvé en Template_bernard_atom.xlsx")
  except:
    tkinter.messagebox.showinfo("Erreur", "Bernard n'a pas réussi le swap. Vérifiez qu'il existe bien un fichier Epure \n Et que ce fichier n'est pas ouvert")

def assign_ca():
  """Transforme le template d'export en template d'import Collective Access"""
  try:
    bernard = openpyxl.load_workbook('Epure.xlsx')
    assign = Workbook()
    bws = bernard.active
    aws = assign.active 
    for i in range(2 ,bws.max_row +1):
      identifier = bws.cell(row=i, column=1).value
      description = bws.cell(row=i, column=2).value
      if bws.cell(row=i, column=3).value != None:
        semicolon_date = re.sub(r',',';', str(bws.cell(row=i, column=3).value)) #Collective access au Carhif ne supporte pas les virgules dans les dates
        dates = semicolon_date
      else:
        dates = bws.cell(row=i, column=3).value
      importance = bws.cell(row=i,column=4).value
      notes = bws.cell(row=i, column= 5).value

      aws.cell(i, 2).value = identifier
      aws.cell(i, 4).value = description
      aws.cell(i, 6).value = dates
      aws.cell(i, 5).value = importance
      aws.cell(i, 19).value = notes
      aws.cell(i, 18).value = 'public_access'

    headers = ['Identifiant de l\'unité','Cote de rangement',
               'Ancienne cote de rangement','Description de l\'unité', 
               'Importance matérielle','Dates', 'ID fonds parent', 'Nom fonds parent',
               'ID fonds parent','ID classement 1','Titre classement 1',
               'ID classement 2','Titre classement 2','ID classement 3',
               'Titre classement 3', 'Nom producteur', 'Type producteur','Visibilité', 'Note']
    r = 1
    for names in headers:
      aws.cell(1, r).value = names
      r+=1

    assign.save('Template_bernard_ca.xlsx')
    tkinter.messagebox.showinfo("Success", "Bernard a tout swappé en template Collective access \n Le fichier a été sauvé en Template_bernard_ca.xlsx")
  except:
    tkinter.messagebox.showinfo("Erreur", "Bernard n'a pas réussi le swap. Vérifiez qu'il existe bien un fichier Epure \n Et que ce fichier n'est pas ouvert")    

def regexMachina():
  try:
    myFile = xlsxwriter.Workbook('Bernard_regex.xlsx')
    worksheet = myFile.add_worksheet()
    headings = ['Numéro', 'Titre unité', 'Dates', 'Importance', 'Notes']
    worksheet.write_row('A1', headings)
    ident = "([1-9]\.|[1-9]\’\.|[1-9][0-9]\.|[1-9][0-9]\’\.|[1-9][0-9][0-9]\.|[1-9][0-9][0-9]\’\.|[0-9]\-[0-9]\.|[0-9]\-[0-9]\s:|[0-9]\-[0-9][0-9]\.|[0-9]\-[0-9][0-9]\.|[1-9][0-9]{3}\.|[1-9][0-9]{3}\’\.)"
    t = "[0-9]\.\\t([A-Za-z\s\'\ô\ö\ç\â\à\ä\ê\é\è\ë\ï\ü\-\—\’\'\(\)\,\“\”\«\»\;\:\\\\&\[\]\°\?\.]+)\\t" #|[0-9]\\t([A-Za-z\s\'\ô\ö\ç\â\à\ä\ê\é\è\ë\ï\ü\-\—\’\'\(\)\,\“\”\«\»\;\:\\\\&\[\]\°\?\.]+)\\t|[0-9]\\t([A-Za-z\s\'\ô\ö\ç\â\à\ä\ê\é\è\ë\ï\ü\-\—\’\'\(\)\,\“\”\«\»\;\:\\\\&\[\]\°\?\.]+)"
    d = "(?<![a-z]\s|:\s|\S0)[1-9][7-9][0-9]{2}-[1-9][7-9][0-9]{2}|(?<![a-z]\s|\,\s)[1-9][7-9][0-9]{2}|(?<![a-z]\s|\,\s)20[0-9]{2}|.\.d\.|sd\.|sans date|\[[1-9][7-9][0-9]{2}-[1-9][7-9][0-9]{2}\]|[ca\.[1-9][7-9][0-9]{2}]"
    imp = "1 p...e|1 p...es|[0-9] p...es|[1-9][0-9] p...es|1 ch....e|1 ch....es|[0-9] ch....es|1 li...e|1 li...es|[0-9] li....s|1 vol..e|[1-9] vol..es"
    row= 0
    row1= 0
    row2= 0
    row3= 0
    col= 0
    rexMachina = docx.Document(c_file)
    text = []
    z =[] 
    for paragraph in rexMachina.paragraphs:
      text.append(paragraph.text)
    print(text)
    x = ' '.join(text)
    print(x)
    for item in re.findall(ident, str(x)):
      #suppression des points initiaux
      cote_light = re.sub("\.", "", str(item))
      #rajout des points pour les mentions bis et '
      cote_d = re.sub("\’|bis",".1", str(cote_light))
      z.append(cote_d)
     
    #Tri + Suppression des doublons pour les identifiants
    z.sort(key=float)
    ident_list = list(dict.fromkeys(z))
   
    for item in ident_list:
      row+=1
      worksheet.write(row, col, str(item))

    for item in re.findall(t, str(x)):
      title_light = re.sub("n\.d\.", "", str(item))
      row1+=1
      worksheet.write(row1, col+1, str(title_light))
     
    for item in re.findall(d, str(x)):
      row2+=1
      worksheet.write(row2, col+2, str(item))
     
    for item in re.findall(imp, str(x)):
      row3+=1
      worksheet.write(row3, col+3, str(item))

    myFile.close()
    tkinter.messagebox.showinfo("Succès", "Le regex a fonctionné ! \nVous pouvez voir le résultat dans Bernard_regex")
  except:
    tkinter.messagebox.showinfo("Echec", "Quelque chose n'a pas fontionné lors du regex\nIl se peut que le fichier Bernard_regex soit déjà ouvert ou que vous n'ayez pas choisi de fichier")
 
def connector():
  return True
